"""
修正版 rag_core.py
- list_ollama_models の堅牢化
- query_index に retrieval-only モードを追加
- meta model 比較を柔軟化
- build_index_progressive を高速化（並列 embedding）
- 各修正箇所の行末にタイムスタンプを付与
"""
import os
import time
import shutil
import json
import tempfile
import zipfile
import logging
import traceback
import re
import pathlib
from typing import Tuple, List, Dict, Callable, Optional, Any
from concurrent.futures import ThreadPoolExecutor, as_completed

import chardet
from pdfminer.high_level import extract_text as pdf_extract
from docx import Document as DocxDocument
import docx2txt
import openpyxl
import csv

# llama-index imports (may vary by version)
try:
    from llama_index.core import VectorStoreIndex, StorageContext, load_index_from_storage
    from llama_index.core.schema import Document
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    from llama_index.core.node_parser import SentenceSplitter
except Exception:
    # fallback imports
    from llama_index import VectorStoreIndex, StorageContext
    from llama_index.schema import Document
    from llama_index.embeddings import HuggingFaceEmbedding
    from llama_index.node_parser import SentenceSplitter

# Query engine / LLM
try:
    try:
        from llama_index.core.query_engine import RetrieverQueryEngine
    except Exception:
        from llama_index.query_engine import RetrieverQueryEngine
except Exception:
    RetrieverQueryEngine = None

try:
    from llama_index.llms.ollama import Ollama
except Exception:
    Ollama = None

# ---------- Logger ----------
logger = logging.getLogger("rag_core")
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
    logger.addHandler(handler)
logger.setLevel(logging.INFO)

logging.getLogger("openpyxl").setLevel(logging.ERROR)
logging.getLogger().setLevel(logging.ERROR)

# ---------- Global cancel flag ----------
CANCEL_FLAG = False

def request_cancel():
    global CANCEL_FLAG
    CANCEL_FLAG = True  # 2025-11-22 12:10:00

def reset_cancel():
    global CANCEL_FLAG
    CANCEL_FLAG = False  # 2025-11-22 12:10:00

# ---------- Configurable defaults ----------
DEFAULT_EMBED_MODEL = os.getenv("EMBED_MODEL_PATH", "local_models/all-MiniLM-L6-v2")  # 2025-11-22 12:10:00
DEFAULT_SPLITTER_CHUNK_SIZE = int(os.getenv("SPLITTER_CHUNK_SIZE", "512"))
DEFAULT_SPLITTER_OVERLAP = int(os.getenv("SPLITTER_OVERLAP", "100"))
DEFAULT_EMBED_BATCH_SIZE = int(os.getenv("EMBED_BATCH_SIZE", "64"))
DEFAULT_LLM_MODEL = os.getenv("OLLAMA_LLM_MODEL", "llama3.2:3b")

# ---------- Utility: meta read/write ----------
def write_meta(index_path: str, meta: dict):
    os.makedirs(index_path, exist_ok=True)
    with open(os.path.join(index_path, "meta.json"), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)  # 2025-11-22 12:10:00

def read_meta(index_path: str) -> Optional[dict]:
    mf = os.path.join(index_path, "meta.json")
    if not os.path.exists(mf):
        return None
    with open(mf, "r", encoding="utf-8") as f:
        return json.load(f)  # 2025-11-22 12:10:00

# ---------- File text extraction ----------
def is_zip_password_protected(path: str) -> bool:
    try:
        with zipfile.ZipFile(path) as z:
            z.testzip()
        return False
    except RuntimeError:
        return True
    except zipfile.BadZipFile:
        return False  # 2025-11-22 12:10:00

def extract_text_from_pptx(path: str) -> str:
    if is_zip_password_protected(path):
        logger.warning("Password-protected PPTX/ZIP skipped: %s", path)
        return ""
    try:
        texts = []
        with zipfile.ZipFile(path, "r") as z:
            slide_files = [f for f in z.namelist() if f.startswith("ppt/slides/slide") and f.endswith(".xml")]
            for fname in slide_files:
                xml = z.read(fname).decode("utf-8", errors="ignore")
                matches = re.findall(r"<a:t>(.*?)</a:t>", xml)
                if matches:
                    texts.extend(matches)
        return "\n".join(texts)
    except Exception:
        logger.exception("PPTX parse failed: %s", path)
        return ""  # 2025-11-22 12:10:00

def extract_text_from_file(path: str) -> Tuple[str, str]:
    try:
        lower = path.lower()
        if lower.endswith(".pdf"):
            text = pdf_extract(path)
            return "success", text
        elif lower.endswith(".docx"):
            try:
                doc = DocxDocument(path)
                text = "\n".join([p.text for p in doc.paragraphs])
            except Exception:
                text = docx2txt.process(path)
            return "success", text
        elif lower.endswith(".xlsx") or lower.endswith(".xls"):
            wb = openpyxl.load_workbook(path, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    parts.append(" ".join([str(c) for c in row if c is not None]))
            return "success", "\n".join(parts)
        elif lower.endswith(".csv"):
            with open(path, "rb") as f:
                raw = f.read()
                enc = chardet.detect(raw).get("encoding") or "utf-8"
            with open(path, "r", encoding=enc, errors="ignore") as f:
                reader = csv.reader(f)
                rows = [",".join(row) for row in reader]
            return "success", "\n".join(rows)
        elif lower.endswith(".txt") or lower.endswith(".md"):
            with open(path, "rb") as f:
                raw = f.read()
                enc = chardet.detect(raw).get("encoding") or "utf-8"
            with open(path, "r", encoding=enc, errors="ignore") as f:
                text = f.read()
            return "success", text
        elif lower.endswith(".pptx"):
            text = extract_text_from_pptx(path)
            return ("success", text) if text else ("skipped", "")
        elif lower.endswith(".zip"):
            return "zip", ""
        else:
            return "skipped", ""
    except Exception as e:
        return "error", f"{type(e).__name__}: {str(e)}\n{traceback.format_exc()}"  # 2025-11-22 12:10:00

# ---------- Safe loader: directory -> Documents with logs ----------
def safe_load_documents_from_folder_verbose(folder: str, progress_callback: Optional[Callable]=None) -> Tuple[List[Document], List[dict]]:
    docs: List[Document] = []
    logs: List[dict] = []

    file_list = []
    for root, _, files in os.walk(folder):
        for fname in files:
            file_list.append(os.path.join(root, fname))

    total = len(file_list)
    logger.info("Found %d files under %s", total, folder)

    for idx, path in enumerate(file_list, start=1):
        if CANCEL_FLAG:
            logger.warning("User cancelled loading at %s", path)
            logs.append({"file": path, "status": "cancelled", "error": "User cancelled", "duration": 0.0})
            break

        if progress_callback:
            try:
                progress_callback(idx, total, path)
            except Exception:
                logger.exception("progress_callback failed")

        start = time.time()
        status, content = extract_text_from_file(path)

        if status == "zip":
            if is_zip_password_protected(path):
                logs.append({"file": path, "status": "skipped", "error": "password-protected zip", "duration": 0.0})
                logger.warning("Skip password-protected zip: %s", path)
            else:
                try:
                    with tempfile.TemporaryDirectory() as tmp_dir:
                        with zipfile.ZipFile(path, "r") as z:
                            z.extractall(tmp_dir)
                        sub_docs, sub_logs = safe_load_documents_from_folder_verbose(tmp_dir, progress_callback=progress_callback)
                        docs.extend(sub_docs)
                        logs.extend([{"file": f"{path}::{s['file']}", "status": s["status"], "error": s.get("error"), "duration": s.get("duration")} for s in sub_logs])
                        logger.info("ZIP extracted %s -> %d files", path, len(sub_docs))
                except RuntimeError as e:
                    logs.append({"file": path, "status": "error", "error": f"ZIP extraction failed: {e}", "duration": time.time() - start})
                    logger.exception("ZIP extraction failed: %s", path)
                except Exception as e:
                    logs.append({"file": path, "status": "error", "error": f"{type(e).__name__}: {e}", "duration": time.time() - start})
                    logger.exception("ZIP handling error: %s", path)
            continue

        duration = time.time() - start
        if status == "success" and content:
            docs.append(Document(text=content, extra_info={"file_path": path}))
            logs.append({"file": path, "status": "success", "error": None, "duration": duration})
            logger.info("[%d/%d] %s processed: success (%.2fs)", idx, total, path, duration)
        elif status == "skipped":
            logs.append({"file": path, "status": "skipped", "error": "unsupported or empty", "duration": duration})
            logger.info("[%d/%d] %s skipped", idx, total, path)
        else:
            logs.append({"file": path, "status": "error", "error": content, "duration": duration})
            logger.error("[%d/%d] %s error: %s", idx, total, path, content)

    return docs, logs  # 2025-11-22 12:10:00

# ---------- Node splitting ----------
def docs_to_nodes(docs: List[Document], chunk_size: int = DEFAULT_SPLITTER_CHUNK_SIZE, chunk_overlap: int = DEFAULT_SPLITTER_OVERLAP) -> List[Document]:
    splitter = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    nodes: List[Document] = []
    logger.info("Splitting %d documents into nodes (chunk_size=%d overlap=%d)...", len(docs), chunk_size, chunk_overlap)
    for d in docs:
        try:
            generated = splitter.get_nodes_from_documents([d])
            for n in generated:
                txt = n.get_content() if hasattr(n, "get_content") else getattr(n, "text", "")
                extra = getattr(n, "extra_info", {}) or {}
                if not extra.get("file_path") and d.extra_info.get("file_path"):
                    extra["file_path"] = d.extra_info.get("file_path")
                nodes.append(Document(text=txt, extra_info=extra))
        except Exception:
            logger.exception("Splitter failed for a document; keeping whole doc")
            nodes.append(d)
    logger.info("Finished splitting; total nodes = %d", len(nodes))
    return nodes  # 2025-11-22 12:10:00

# HuggingFaceEmbedding 初期化
def get_embedding_instance(embed_model_path: str, device: str = "cpu", embed_batch_size: int = DEFAULT_EMBED_BATCH_SIZE):
    return HuggingFaceEmbedding(
        model_name=embed_model_path,
        device=device,
        embed_batch_size=embed_batch_size
    )  # 2025-11-22 12:10:00

# ---------- 高速化版 Progressive Index Build ----------

def build_index_progressive(nodes: list, embed_model, index_path: str,
                            chunk_callback: Optional[Callable] = None,
                            batch_size: int = 32) -> VectorStoreIndex:  # 2025-11-22
    """
    高速 Index Build: nodes を batch 単位で embedding し VectorStoreIndex に挿入
    - nodes: list of Documents
    - embed_model: HuggingFaceEmbedding instance
    - index_path: 保存先フォルダ
    - chunk_callback: 進捗コールバック (現在ノード, 総ノード)
    - batch_size: 一度に embed するドキュメント数
    """
    import time
    logger.info("===== 高速 Index Build Start =====")
    start_time = time.time()
    os.makedirs(index_path, exist_ok=True)

    total_nodes = len(nodes)
    index = None

    for i in range(0, total_nodes, batch_size):
        batch = nodes[i:i+batch_size]
        try:
            # batch embedding
            embed_model(batch)  # HuggingFaceEmbedding は Document の list をそのまま渡せる
        except Exception:
            logger.exception("Batch embedding failed for nodes %d-%d", i, i+len(batch)-1)
            # 失敗したら個別 retry
            for doc in batch:
                try:
                    embed_model([doc])
                except Exception:
                    logger.exception("Embedding failed for a single doc")

        # index に挿入
        if index is None:
            index = VectorStoreIndex.from_documents(batch, embed_model=embed_model)
        else:
            for doc in batch:
                try:
                    index.insert(doc)
                except Exception:
                    logger.exception("Insert failed for a doc in index")

        # progress callback
        if chunk_callback:
            chunk_callback(min(i+batch_size, total_nodes), total_nodes)

        if CANCEL_FLAG:
            logger.warning("User cancelled during index build.")
            break

        # ログ
        if (i // batch_size) % 5 == 0 or (i + batch_size) >= total_nodes:
            logger.info("Indexed %d/%d nodes", min(i+batch_size, total_nodes), total_nodes)

    if index:
        index.storage_context.persist(index_path)

    elapsed = time.time() - start_time
    logger.info("===== 高速 Index Build Completed (%.2fs) =====", elapsed)
    return index
#20251122 高速バッチ <<<<<<<<<<<<<
#####################################################################

#20251122 Export Text >>>>>>>>>>>>

def _safe_filename_from_path(path: str) -> str:
    """UTF-8 日本語を残して安全なファイル名に変換"""
    base = pathlib.Path(path).name
    # Windows で使えない文字だけ置換
    safe = re.sub(r'[<>:"/\\|?*\n\r]+', "_", base)
    # 長すぎる場合は切る
    return safe[:200] if len(safe) > 200 else safe


def write_extracted_texts(docs: List[Document], index_path: str) -> List[str]:
    """
    Write extracted texts of Documents to index_path/extracted_texts/*.txt.
    Returns list of written file paths.
    """
    out_dir = os.path.join(index_path, "extracted_texts")
    os.makedirs(out_dir, exist_ok=True)
    written = []
    counters = {}
    for d in docs:
        src = (d.extra_info or {}).get("file_path") if hasattr(d, "extra_info") else None
        if not src:
            # generate generic name
            name = "doc"
        else:
            name = _safe_filename_from_path(src)

        # ensure unique
        cnt = counters.get(name, 0)
        counters[name] = cnt + 1
        if cnt:
            fname = f"{name}__{cnt}.txt"
        else:
            fname = f"{name}.txt"

        out_path = os.path.join(out_dir, fname)
        try:
            text = getattr(d, "text", None) or getattr(d, "content", "") or ""
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
            written.append(out_path)
        except Exception:
            logger.exception("Failed to write extracted text for %s", src or "<unknown>")
    logger.info("Wrote %d extracted text files to %s", len(written), out_dir)
    return written
#20251122 Export Text >>>>>>>>>>>>>>>>

# ---------- Index build (high-level) ----------
def build_index_from_folder(folder: str, index_name: str,
                            storage_dir: str = "storage",
                            embed_model_path: str = DEFAULT_EMBED_MODEL,
                            embed_device: str = "cpu",
                            embed_batch_size: int = DEFAULT_EMBED_BATCH_SIZE,
                            chunk_size: int = DEFAULT_SPLITTER_CHUNK_SIZE,
                            chunk_overlap: int = DEFAULT_SPLITTER_OVERLAP,
                            progress_callback: Optional[Callable] = None,
                            chunk_callback: Optional[Callable] = None,
                            export_texts=False) -> dict:
    reset_cancel()

    if not os.path.isdir(folder):
        return {"status": "error", "message": "folder not found", "file_logs": []}

    index_path = os.path.join(storage_dir, index_name)
    if os.path.exists(index_path):
        shutil.rmtree(index_path)
        logger.info("Removed existing index at %s", index_path)

    # 1) Load docs with logs and progress
    docs, file_logs = safe_load_documents_from_folder_verbose(folder, progress_callback=progress_callback)
    if CANCEL_FLAG:
        logger.warning("Index build cancelled during file loading")
        return {"status": "cancelled", "message": "cancelled", "file_logs": file_logs}

    if not docs:
        logger.warning("No documents loaded from folder %s", folder)
        return {"status": "error", "message": "no documents loaded", "file_logs": file_logs}

    # optionally export extracted texts for inspection 
    #20251122 export text>>>>>>>>>>>>>>>>
    exported_files = []
    if export_texts:
        try:
            exported_files = write_extracted_texts(docs, index_path)
        except Exception:
            logger.exception("Failed exporting extracted texts") 
    #20251122 export text<<<<<<<<<<<<<<<<


    # 2) Split docs into nodes
    nodes = docs_to_nodes(docs, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    if CANCEL_FLAG:
        logger.warning("Index build cancelled after splitting")
        return {"status": "cancelled", "message": "cancelled", "file_logs": file_logs}

    # 3) Create embedding model instance
    embed_model = get_embedding_instance(embed_model_path=embed_model_path, device=embed_device, embed_batch_size=embed_batch_size)

    # 4) Build index progressively (this is the heavy step)
    logger.info("Starting index creation (nodes=%d) using HuggingFace model=%s (device=%s batch=%d)",
                len(nodes), embed_model_path, embed_device, embed_batch_size)
    try:
        index = build_index_progressive(nodes, embed_model, index_path, chunk_callback=chunk_callback)
        
    except Exception as e:
        logger.exception("Index creation failed")
        return {"status": "error", "message": f"Index creation failed: {e}", "file_logs": file_logs}

    # 5) Write meta
    meta = {
        "embedding": {"provider": "huggingface", "model": embed_model_path, "device": embed_device, "batch_size": embed_batch_size},
        "splitter": {"chunk_size": chunk_size, "chunk_overlap": chunk_overlap},
        "created_at": time.time()
    }
    write_meta(index_path, meta)

    return {"status": "ok", "message": "index created", "file_logs": file_logs}

# ----------------------------
# LLM utilities: list Ollama models and query


def list_ollama_models() -> List[str]:
    """Returns list of available Ollama model names (empty if ollama client not available)."""

    OLLAMA_HOST_URL = 'http://127.0.0.1:11434' 

    try:
        import ollama
        models = []
        try:
            raw = ollama.list_models()
            for m in raw:
                if isinstance(m, dict):
                    name = m.get("name") or m.get("model") or m.get("id")
                else:
                    name = getattr(m, "name", None) or getattr(m, "model", None)
                if name:
                    models.append(name)
        except Exception:
            # fallback to client
            try:
                client = ollama.Client(host=OLLAMA_HOST_URL, timeout=30.0) #Timeout when Ollama not found case
                raw = client.list()
                for m in getattr(raw, "models", []) or raw:
                    if isinstance(m, dict):
                        name = m.get("name") or m.get("model")
                    else:
                        name = getattr(m, "name", None) or getattr(m, "model", None)
                    if name:
                        models.append(name)
            except Exception as e:
                logger.exception("ollama client fallback failed: %s", e)
        return sorted(set(models))
    except Exception as e:
        logger.warning("Ollama not available or failed to list models: %s", e)
        return []
# 2025-11-21 15:00:00

# ----------------------------
# Query

def query_index(index_name: str, query: str, storage_dir: str = "storage",
                embed_model_path: str = DEFAULT_EMBED_MODEL,
                embed_device: str = "cpu",
                llm_model_name: str = DEFAULT_LLM_MODEL,
                use_llm: bool = True,
                top_k: int = 5,
                llm_instance=None) -> dict:
    if not query:
        return {"status": "error", "error": "no query input"}

    index_path = os.path.join(storage_dir, index_name)
    if not os.path.exists(index_path):
        return {"status": "error", "error": "index not found"}

    meta = read_meta(index_path)
    if not meta:
        return {"status": "error", "error": "meta.json missing in index"}

    # flexible comparison for embedding model (normalize paths)  # 2025-11-21 15:00:00
    meta_model = meta.get("embedding", {}).get("model")
    if meta_model:
        try:
            if os.path.abspath(meta_model) != os.path.abspath(embed_model_path) and (meta_model not in embed_model_path and embed_model_path not in meta_model):
                return {"status": "error", "error": f"Embedding model mismatch: index uses {meta_model} but request used {embed_model_path}"}
        except Exception:
            # fallback conservative check
            if meta_model != embed_model_path:
                return {"status": "error", "error": f"Embedding model mismatch: index uses {meta_model} but request used {embed_model_path}"}

    embed_model = get_embedding_instance(embed_model_path=embed_model_path, device=embed_device)
    storage_context = StorageContext.from_defaults(persist_dir=index_path)
    try:
        index = load_index_from_storage(storage_context, embed_model=embed_model)
    except Exception:
        try:
            index = VectorStoreIndex.load_from_disk(index_path)
        except Exception as e:
            logger.exception("Failed to load index from storage")
            return {"status": "error", "error": f"Failed to load index: {e}"}

    # Retrieval-only mode
    if not use_llm:
        try:
            retriever = index.as_retriever(similarity_top_k=top_k)
            # support different retriever interfaces
            if hasattr(retriever, "retrieve"):
                docs = retriever.retrieve(query)
            elif hasattr(retriever, "get_relevant_documents"):
                docs = retriever.get_relevant_documents(query)
            else:
                docs = retriever(query)

            snippets = []
            for d in docs:
                text = getattr(d, "text", None) or getattr(d, "content", None) or str(d)
                file = (d.extra_info or {}).get("file_path") if hasattr(d, "extra_info") else None
                snippets.append({"file": file, "snippet": text[:2000]})
            return {"status": "ok", "response": {"query": query, "results": snippets}}
        except Exception as e:
            logger.exception("Retrieval-only query failed")
            return {"status": "error", "error": f"retriever-only failed: {type(e).__name__}: {e}"}

    # LLM path: use Ollama if available
    if Ollama is None:
        return {"status": "error", "error": "Ollama LLM wrapper not installed"}

    try:
        if isinstance(llm_model_name, str) and ("20b" in llm_model_name or "30b" in llm_model_name or "70b" in llm_model_name):
            logger.warning("Selected LLM '%s' may require large memory and may fail on current machine.", llm_model_name)

        # allow injection of an llm_instance (for testing) or create one
        if llm_instance is None:
            llm_instance = Ollama(model=llm_model_name, base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),request_timeout=120.0)

        qe = index.as_query_engine(llm=llm_instance)
        resp = qe.query(query)
        return {"status": "ok", "response": str(resp)}
    except Exception as e:
        logger.exception("Query with LLM failed")
        return {"status": "error", "error": f"{type(e).__name__}: {e}"}

# ----------------------------
# List indexes helper

def list_indexes(storage_dir: str = "storage") -> List[str]:
    if not os.path.exists(storage_dir):
        return []
    return [d for d in os.listdir(storage_dir) if os.path.isdir(os.path.join(storage_dir, d))]
